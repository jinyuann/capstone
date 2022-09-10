# hello.py
import os
from flask import Flask, render_template, request,send_file
import docx


app = Flask(__name__)

data = []
content = None
SDSinsertion = False



@app.route('/')
def hello():
    return render_template('index.html')

@app.route('/uploader',methods = ['POST', 'GET'])
def login():
   if request.method == 'POST':
      data.clear()
      file = request.files['file']
      print (file) #File type
      doc = docx.Document(file)
      table_count = len(doc.tables)
      for index in range (table_count):

        table = doc.tables[index]

        for i, row in enumerate(table.rows):
          text = (cell.text for cell in row.cells)
          if i == 0:
              keys = tuple(text)
              continue
          row_data = tuple(text)
          data.append(row_data)
      for row_data in range(len(data)):
          print (data[row_data])

      if data:
          SDSinsertion = True

      return render_template('index.html',SDSinsertion=SDSinsertion)

@app.route('/Converter', methods=['POST', 'GET'])
def CobolJCL():
    try:
       if request.method == 'POST':
           data.clear()
           file2 = request.files['file']
           if not file2:
               SDSinsertion = True
               return render_template('index.html',SDSinsertion=SDSinsertion)
           print(file2)  # File type
           doc = docx.Document(file2)
           table_count = len(doc.tables)
           for index in range(table_count):

               table = doc.tables[index]

               for i, row in enumerate(table.rows):
                   text = (cell.text for cell in row.cells)
                   if i == 0:
                       keys = tuple(text)
                       continue
                   row_data = tuple(text)
                   data.append(row_data)
           for row_data in range(len(data)):
               print(data[row_data])

           Environment = request.form['environment']
           Spooling = request.form['spooling']
           print(Environment)
           print(Spooling)
           content = request.files['file2'].readlines()
           if not content:
               CobolJCLinsertion = True
               return render_template('index.html',CobolJCLinsertion=CobolJCLinsertion)
           print(content)  # File content
           headercount=0
           detailcount = 0
           summaryfile=''
           for x in content:
               headercount +=1
               line = (x.decode("utf-8"))  # Remove binary
               linew = line.replace('\n', '')  # Strip dirty data
               fullline = linew.replace('\r','')  # Strip dirty data, lineww is the string value of each line in the text fi
               #print (fullline)
               if headercount==1:
                   JCLname = fullline.split() #Split into word by word
                   JCLname = JCLname[0].replace('//', '') #remove // if any
                   JCLname = JCLname +'@'
                   NewJCLFile =  open(JCLname, "w")
               #USERPROC TO ADD IN ENVIRONMENT
               UserProc = fullline.split()  # Split into word by word
               if 1 < len (UserProc):
                   #print (MSGLevel[1]) finding MSGLEVEL
                   if "MSGLEVEL" in UserProc[1]:
                       NewJCLFile.write("//     MSGLEVEL=(1,1),REGION=0M,NOTIFY=&SYSUID,COND=(4,LT)" + '\n')
                       continue
               UserProc = UserProc[0].replace('//', '')  # remove // if any
               if UserProc == "USERPROC":
                   if Environment =="OB1_DEV":
                    NewJCLFile.write("//USERPROC JCLLIB ORDER=PRD6.ONB.DEV.SRC.JCLPROC" + '\n')
                    continue
                   elif Environment == "OB1_UAT":
                    NewJCLFile.write("//USERPROC JCLLIB ORDER=PRD6.ONB.UAT.SRC.JCLPROC" + '\n')
                    continue
                   elif  Environment =="OB2_DEV":
                    NewJCLFile.write("//USERPROC JCLLIB ORDER=PRD6.ONB2.DEV.SRC.JCLPROC" + '\n')
                    continue
                   elif Environment == "OB2_UAT":
                    NewJCLFile.write("//USERPROC JCLLIB ORDER=PRD6.ONB2.UAT.SRC.JCLPROC" + '\n')
                    continue
                   elif Environment == "PROD":
                    NewJCLFile.write("//USERPROC JCLLIB ORDER=PRD3.RCS.JCLPROC.SRCLIB" + '\n')
                    continue
               if fullline.upper() =="REMOVE1":
                   detailcount =1
                   NewJCLFile.write("//*---------------------------------------------------------------------------*" + '\n')
                   NewJCLFile.write("//*                SET BATCH CLIENT ARGUMENT                                  *" + '\n')
                   NewJCLFile.write("//*---------------------------------------------------------------------------*" + '\n')
                   NewJCLFile.write("//STEPXX   EXEC PGM=IEBGENER" + '\n')
                   NewJCLFile.write("//SYSIN    DD DUMMY" + '\n')
                   NewJCLFile.write("//SYSPRINT DD SYSOUT=*" + '\n')
                   NewJCLFile.write("//SYSUT2   DD DSN=&&JAVBTHC$,DISP=(NEW,PASS,DELETE)," + '\n')
                   NewJCLFile.write("//         DCB=(LRECL=80,RECFM=FB,BLKSIZE=8000)," + '\n')
                   NewJCLFile.write("//         SPACE=(TRK,(10,10),RLSE)" + '\n')
                   NewJCLFile.write("//SYSUT1   DD *" + '\n')
                   for row_data in range(len(data)):
                       spacecount=0
                       spaces=""
                       # Adding missing spaces for the parameters
                       for i in range (12):
                           if len(data[row_data][1]) < i:
                             spacecount = spacecount + 1
                       # Adding missing spaces for the parameters
                       for i in range (spacecount):
                           spaces = spaces +" "
                       parameterspace = data[row_data][1] + spaces
                       NewJCLFile.write("'" +parameterspace + ' >' + data[row_data][2] + "'" + '\n')
                       if "SUMMARY.SD" in data[row_data][2]:
                           summaryfile = data[row_data][2]
                   NewJCLFile.write("/*" + '\n')
                   NewJCLFile.write("//*---------------------------------------------------------------------------*" + '\n')
                   NewJCLFile.write("//*                TRIGGER JAVA BATCH CLIENT CLASS                            *" + '\n')
                   NewJCLFile.write("//*---------------------------------------------------------------------------*" + '\n')
                   NewJCLFile.write("//STEPXX   EXEC PROC=ZOSJVBAT" + '\n')
                   if Spooling == "yes":
                       NewJCLFile.write("//*************SPOOL DISPLAY DATASET TO CONSOLE**********************" + '\n')
                       NewJCLFile.write("//STEPXX EXEC PGM=IEBGENER" + '\n')
                       NewJCLFile.write("//SYSIN DD DUMMY" + '\n')
                       NewJCLFile.write("//SYSPRINT DD SYSOUT=X" + '\n')
                       NewJCLFile.write("//SYSOUT DD SYSOUT=*" + '\n')
                       NewJCLFile.write("//SYSUT1 DD DISP=SHR,DSN=" + summaryfile + '\n')
                       NewJCLFile.write("//SYSUT2 DD SYSOUT=*" + '\n')
                       NewJCLFile.write("//*" + '\n')
                       NewJCLFile.write("//******************************************************************" + '\n')

               elif (fullline.upper() =="REMOVE2"):
                   detailcount = 0
                   continue
               if detailcount ==0:
                   NewJCLFile.write(fullline+'\n')
           NewJCLFile.close()
           data.clear()

           print (os.path.abspath(JCLname))

           success=True
           return render_template('index.html',success=success,JCLname=JCLname)
    except Exception as e:
        print(e)
        Error = True
        return render_template('index.html',Error=Error)







@app.route('/downloads', methods=['POST', 'GET'])
def downloads():
           name = request.args.get("id")
           print (name)
           return send_file(
               os.path.abspath(name),
               mimetype='text',
               download_name=name,
               as_attachment=True
           )


if __name__ == '__main__':
    app.run(host='0.0.0.0', debug=True)
